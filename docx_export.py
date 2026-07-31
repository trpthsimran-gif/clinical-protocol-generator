"""
docx_export.py

Converts the AI-generated protocol draft (markdown-ish text with
## headings, **bold**, and PMID links) into a clean, downloadable
Word (.docx) document using python-docx. This is an alternative to
pdf_export.py for people who need an editable document (e.g. to add
their own institution's letterhead, or to track changes in Word).

Handles the same markdown patterns as pdf_export.py:
    - "## Heading" lines      -> Word heading styles
    - "**bold text**"         -> bold inline runs
    - "[label](url)" links    -> real clickable Word hyperlinks
    - "- item" / "* item"     -> bullet list items
"""

import re
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_hyperlink(paragraph, url: str, text: str):
    """
    python-docx has no built-in hyperlink support, so this manually builds
    the required XML (a well-known recipe for this library).
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1a73e8")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_runs(paragraph, text: str):
    """
    Parse a single line for **bold** spans and [label](url) links,
    adding each piece as the correct kind of run in order.
    """
    # Split on links first, keeping the delimiters so we can process in order
    link_pattern = re.compile(r"\[([^\]]+)\]\((https?://[^\)]+)\)")
    pos = 0
    for match in link_pattern.finditer(text):
        # Add any plain/bold text before this link
        _add_bold_runs(paragraph, text[pos:match.start()])
        _add_hyperlink(paragraph, match.group(2), match.group(1))
        pos = match.end()
    # Add any remaining text after the last link
    _add_bold_runs(paragraph, text[pos:])


def _add_bold_runs(paragraph, text: str):
    """Split text on **bold** markers and add plain/bold runs accordingly."""
    if not text:
        return
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # odd-indexed parts were inside **...**
            run.bold = True


def generate_protocol_docx(topic: str, draft_markdown: str) -> bytes:
    """
    Build a Word document from the AI-generated draft and return it as
    raw bytes, ready to hand to Streamlit's download button.
    """
    doc = Document()

    # --- Header ---
    title = doc.add_heading("Clinical Protocol Draft", level=0)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Topic: {topic}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    meta2 = doc.add_paragraph()
    meta2_run = meta2.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')} — "
        "AI-generated draft. Requires clinician review before use."
    )
    meta2_run.font.size = Pt(9)
    meta2_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # --- Body: parse the markdown-ish draft line by line ---
    for raw_line in draft_markdown.splitlines():
        line = raw_line.strip()

        if not line:
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)

    # --- Footer disclaimer (always included, regardless of what GPT wrote) ---
    doc.add_paragraph()
    sep = doc.add_paragraph("_" * 60)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "This is an AI-generated draft grounded in retrieved literature. "
        "It is not validated for clinical use. A qualified clinician must "
        "review and approve this protocol before any real-world application."
    )
    disclaimer_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    disclaimer_run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
